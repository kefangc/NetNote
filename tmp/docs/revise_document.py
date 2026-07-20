from __future__ import annotations

from copy import copy
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"F:\projects\软件杯")
SOURCE = ROOT / "docs" / "NetNote-中国软件杯-项目文档（4）.docx"
OUTPUT = ROOT / "docs" / "NetNote-中国软件杯-项目文档（4）-已校准.docx"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def fit_text(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], *, bold: bool = False, color: str = "#17324d"):
    left, top, right, bottom = box
    max_width = max(1, right - left - 16)
    for size in range(max(11, (bottom - top) // 2), 10, -1):
        candidate = font(size, bold)
        if draw.textbbox((0, 0), text, font=candidate)[2] <= max_width:
            break
    else:
        candidate = font(10, bold)
    bbox = draw.textbbox((0, 0), text, font=candidate)
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=candidate, fill=color)


def rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = "#a8bfd4"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=2)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#4a7899"):
    draw.line([start, end], fill=color, width=4)
    x, y = end
    draw.polygon([(x, y), (x - 14, y - 8), (x - 14, y + 8)], fill=color)


def architecture_image(size: tuple[int, int]) -> bytes:
    width, height = size
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    fit_text(draw, "NetNote V1 实际运行架构", (0, 18, width, int(height * 0.12)), bold=True, color="#102a43")

    rows = [
        ("浏览器端：Next.js + React + TypeScript", "#e8f2fb"),
        ("FastAPI API：来源、聊天、生成物、学习反馈", "#eaf6f0"),
        ("职责模块：检索 | 辅导 | 资源生成 | 画像 | 安全 | 搜索", "#f4eefb"),
        ("本地持久化：workspace.json + uploads", "#fff4db"),
    ]
    top = int(height * 0.16)
    row_h = int(height * 0.13)
    for index, (label, fill) in enumerate(rows):
        y = top + index * int(height * 0.17)
        box = (int(width * 0.12), y, int(width * 0.88), y + row_h)
        rounded(draw, box, fill)
        fit_text(draw, label, box, bold=index < 2)
        if index:
            arrow(draw, (width // 2, y - int(height * 0.035)), (width // 2, y - 4))

    side_y = int(height * 0.36)
    for label, x, fill in [
        ("网页检索与正文抽取", int(width * 0.02), "#f2f7fb"),
        ("OpenAI-compatible 模型", int(width * 0.74), "#f2f7fb"),
        ("云大学堂转写来源", int(width * 0.02), "#f2f7fb"),
        ("Studio 学习资源", int(width * 0.74), "#f2f7fb"),
    ]:
        y = side_y if "网页" in label or "OpenAI" in label else int(height * 0.62)
        box = (x, y, x + int(width * 0.23), y + int(height * 0.11))
        rounded(draw, box, fill, "#bed0df")
        fit_text(draw, label, box)
    return image_bytes(image, "JPEG")


def workflow_image(size: tuple[int, int]) -> bytes:
    width, height = size
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    fit_text(draw, "NetNote V1 来源处理与学习闭环", (0, 16, width, int(height * 0.12)), bold=True, color="#102a43")
    labels = [
        ("文件 / 网页 / 课堂转写", "#e8f2fb"),
        ("解析、正文抽取、切片", "#eaf6f0"),
        ("JSON 来源库与引用片段", "#fff4db"),
        ("检索问答与 Studio 生成", "#f4eefb"),
        ("测验、闪卡反馈与建议", "#eaf6f0"),
    ]
    margin = int(width * 0.04)
    gap = int(width * 0.018)
    card_w = (width - 2 * margin - 4 * gap) // 5
    y1, y2 = int(height * 0.32), int(height * 0.60)
    for index, (label, fill) in enumerate(labels):
        x = margin + index * (card_w + gap)
        box = (x, y1, x + card_w, y2)
        rounded(draw, box, fill)
        fit_text(draw, label, box, bold=index in {0, 3})
        if index:
            arrow(draw, (x - gap + 2, (y1 + y2) // 2), (x - 4, (y1 + y2) // 2))
    arrow(draw, (int(width * 0.90), int(height * 0.72)), (int(width * 0.14), int(height * 0.72)), "#658c75")
    fit_text(draw, "学习反馈更新画像，并生成下一步学习建议", (int(width * 0.24), int(height * 0.77), int(width * 0.76), int(height * 0.90)), color="#3d6b53")
    return image_bytes(image, "PNG")


def image_bytes(image: Image.Image, fmt: str) -> bytes:
    stream = BytesIO()
    image.save(stream, format=fmt, quality=92 if fmt == "JPEG" else None)
    return stream.getvalue()


def replace_paragraph(paragraph, text: str, style: str | None = None):
    paragraph.text = text
    if style:
        paragraph.style = style
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def replace_table(table, headers: list[str], rows: list[list[str]]):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.columns) < len(headers):
        for row in table.rows:
            row._tr.add_tc()
    data = [headers, *rows]
    for index, values in enumerate(data):
        if index >= len(table.rows):
            table.add_row()
        row = table.rows[index]
        for column, value in enumerate(values):
            cell = row.cells[column]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8)


def insert_image_after(paragraph, image: bytes, width: Mm):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.alignment = 1
    inserted.add_run().add_picture(BytesIO(image), width=width)
    return inserted


def replace_media(docx_path: Path):
    with ZipFile(docx_path, "r") as source:
        contents = {item.filename: source.read(item.filename) for item in source.infolist()}
    replacements = {}
    for name, builder in {
        "word/media/image3.jpeg": architecture_image,
        "word/media/image4.jpeg": architecture_image,
        "word/media/image6.jpeg": workflow_image,
        "word/media/image7.png": workflow_image,
    }.items():
        raw = contents.get(name)
        if not raw:
            continue
        size = Image.open(BytesIO(raw)).size
        replacements[name] = builder(size)
    contents.update(replacements)
    temp_path = ROOT / "tmp" / "docs" / "netnote_calibrated_media.docx"
    with ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for name, data in contents.items():
            target.writestr(name, data)
    temp_path.replace(docx_path)


def main():
    doc = Document(SOURCE)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    paragraphs = doc.paragraphs
    replacements = {
        63: "本说明书面向 NetNote V1 原型，记录已实现的来源导入、轻量检索问答、学习资源生成、学习反馈与课程转写导入能力，并明确列出当前限制和后续演进方向。",
        70: "NetNote 是面向高校计算机网络课程的个性化学习工作台。学生可导入课程资料、网络资料或课堂转写文本，在来源基础上进行问答，并生成学习资源和复习建议。系统以可演示、可扩展的 V1 原型为定位，强调来源依据、学习反馈和清晰的交互闭环。",
        71: "1. 来源学习与问答：学生上传 PDF、DOCX、PPTX、TXT、Markdown 等资料后，系统抽取可读文本并切分为来源片段。提问时，后端从已导入片段中检索相关证据，返回带来源标题和位置的回答；缺少直接来源时，会明确说明回答属于通用解释。",
        72: "2. 网络资料补充：学生在来源面板中输入检索词，系统返回候选网页。导入时依次尝试 Jina Reader、Trafilatura 和 HTML 正文清洗，并在正文获取不完整时保留搜索摘要作为降级结果。网页与本地文件进入同一来源库，后续可用于问答和资源生成。",
        73: "3. 云大学堂课堂转写：系统支持连接云大学堂、查询课程并导入可获取的课堂转写内容。转写片段保留课程、章节、起止时间和视频地址等元数据；命中课堂片段的引用可打开播放器并跳转到对应时间段。",
        74: "4. Studio 学习资源：系统可生成综合概要、闪卡、测验、思维导图、问答卡片、拓展阅读和演示文稿。测验和闪卡支持交互式反馈，演示文稿支持浏览器端预览、PDF 导出和图片版 PPTX 导出。",
        75: "5. 学习反馈与建议：系统根据聊天兴趣、测验正确率和闪卡“会/不会/跳过”记录，更新薄弱点、资源偏好、兴趣与下一步学习建议，帮助学生形成可持续的复习路径。",
        77: "NetNote 的技术亮点在于将来源管理、轻量检索问答、学习资源生成和学习反馈放在同一工作台中实现。后端按检索、辅导、资源生成、画像、安全与网页搜索等职责拆分模块；前端提供稳定的三栏学习界面。系统同时提供网络来源正文抽取降级策略，以及课堂转写片段到视频播放的联动能力。",
        80: "NetNote 前端采用 Next.js 16 App Router、React 19、TypeScript 和 Tailwind CSS，实现来源、聊天与 Studio 三栏工作台。后端采用 FastAPI 和 Pydantic，提供来源导入、网络搜索、流式问答、资源生成、测验提交、闪卡复习和学习画像接口。",
        81: "前端通过 React 状态和 API 调用维护工作区数据，并提供 Markdown 渲染、思维导图缩放、学习资源详情和浏览器端导出等交互能力。",
        82: "后端将来源、消息、生成物、学习画像和学习事件保存为本地 JSON，适合单机演示。模型侧通过 OpenAI-compatible 接口配置服务；未配置模型或调用失败时，系统可使用本地规则与来源片段生成基础演示结果。",
        83: "当前 V1 以本地运行和演示稳定性为优先，暂未引入向量数据库、正式 Embedding 检索、OCR 和显式图编排。上述能力作为后续 V2 的可演进方向。",
        85: "项目内置计算机网络种子知识与可导入的课程资料、网页候选和课堂转写来源。文件导入支持多种常用教学资料格式；网络来源支持正文抽取与摘要降级，便于在演示时覆盖不同类型的学习素材。",
        87: "系统面向课程复习、资料整理、概念问答和课堂内容回看等场景。通过将资料来源、生成式资源与学习反馈集中在同一工作区，可降低在多个工具间切换的成本；实际学习效果仍需在后续真实教学场景中持续评估。",
        91: "静态数据包括课程名称、种子知识和界面配置；动态数据包括来源元数据、来源片段、对话消息、生成物、学习画像和学习事件。V1 通过本地 JSON 文件保存工作区状态，上传文件保存在本地 uploads 目录。",
        93: "V1 使用 Pydantic 数据模型约束来源、片段、消息、生成物和学习画像。下表描述的是运行时对象字段，不是 PostgreSQL 物理表或外键关系。",
        94: "表 3-1 来源与片段运行时模型（Source / SourceChunk）",
        95: "表 3-2 学习画像运行时模型（LearningProfile）",
        97: "围绕课程资料学习，本系统实现以下核心功能，并将需要持续优化的能力明确标记为后续方向：",
        98: "需求一：学习画像随学更新。系统通过聊天关键词、测验结果和闪卡反馈维护知识基础描述、学习目标、薄弱点、资源偏好、正确率、学习节奏、易错模式、兴趣和下一步建议。",
        99: "需求二：按职责拆分的学习资源生成。检索、辅导、资源生成、画像、安全和网页搜索模块协作完成概要、闪卡、测验、思维导图、问答卡片、拓展阅读和演示文稿生成。",
        100: "需求三：学习建议。系统基于薄弱点、测验正确率和学习记录生成下一步复习建议，例如先复习薄弱概念、完成一轮闪卡、再进行测验验证。",
        101: "需求四：网络来源搜索与导入。学生可主动搜索网页候选并选择导入；系统提取正文、记录抓取质量并按统一来源结构保存，供后续检索和生成使用。",
        102: "需求五：云大学堂课堂转写与片段播放。系统可导入课程转写片段及其时间元数据；课堂引用可打开视频或课程页面，并从对应片段开始播放。",
        103: "需求六：Studio 学习资源与演示文稿。Studio 提供概要、闪卡、测验、思维导图、问答卡片、拓展阅读和演示文稿。演示文稿由前端渲染，可导出 PDF 和图片版 PPTX。",
        106: "交互体验：聊天和生成任务均提供加载状态与渐进输出，Markdown、表格和代码块保持可读。性能指标以实际运行环境和模型服务为准，本文不使用未经压测验证的固定时延数据。",
        107: "可靠性：文件解析、网页正文抽取和模型调用均设置降级路径。网页正文不足时可回退到搜索摘要；模型不可用时可使用本地启发式回答或生成策略，避免演示流程因外部服务不可用而中断。",
        110: "4.1 V1 前后端分离架构设计",
        111: "系统由 Next.js 前端工作台和 FastAPI 后端服务组成。前端负责来源、聊天和 Studio 交互；后端负责文件解析、网页抽取、来源切片、检索、问答、资源生成和学习反馈。工作区状态保存到本地 JSON，上传文件保存到本地目录；模型服务通过 OpenAI-compatible 接口按环境配置接入。",
        112: "图 4-1 NetNote V1 实际运行架构图",
        116: "图 4-2 NetNote 三栏学习工作台",
        118: "1. 来源导入：用户上传文件、选择网络候选或导入课堂转写；后端解析可读文本并生成来源片段，随后更新本地工作区状态。",
        119: "2. 来源问答：用户问题进入 Retrieval 模块，系统按关键词、全文、来源质量和元数据匹配相关片段；Tutor 模块将证据片段组装为提示上下文，返回回答和引用信息。",
        120: "3. Studio 与反馈：Resource 模块基于来源生成学习资源；测验和闪卡结果由 Profile 模块汇总，更新薄弱点与下一步建议；Safety 模块对明显不安全内容进行基础拦截。",
        126: "图 4-3 NetNote V1 来源处理与学习反馈闭环",
        129: "系统为网页抽取、模型调用和来源不足等场景提供降级处理。网页正文提取失败时保留搜索摘要并标记质量；没有检索到直接来源时，问答可继续给出通用解释并明确说明；模型服务不可用时使用本地启发式结果完成基本演示。",
        130: "第五章：数据模型与本地持久化设计",
        131: "V1 以本地 JSON 文件作为工作区状态存储，数据模型由 Pydantic 定义。该设计降低了原型部署门槛，但不适合多用户、高并发场景；PostgreSQL 和 pgvector 被保留为后续演进方案。",
        132: "5.1 来源与切片数据模型",
        133: "表 5-1 SourceChunk 运行时字段定义",
        134: "5.2 云大学堂课程转写元数据",
        135: "课堂转写片段在 metadata 中保存课程、章节、起止时间、视频地址和课程页面地址等信息，用于课堂引用和片段播放。",
        136: "表 5-2 Lecture 元数据字段定义",
        138: "5.3 Studio 生成物数据模型",
        139: "表 5-3 Artifact 运行时字段定义",
        140: "5.4 对话消息与引用数据模型",
        141: "表 5-4 Message / Citation 运行时字段定义",
        143: "第六章：详细设计与功能实现",
        144: "6.1 文件解析、切片与来源管理",
        145: "后端 parsers.py 根据文件扩展名调用相应解析器：PDF 使用 pdfplumber，DOCX 使用 python-docx，PPTX 使用 python-pptx，TXT 和 Markdown 直接读取文本。图片来源当前只登记 OCR 管线入口，未将图片 OCR 作为 V1 已实现能力。解析后的文本按长度切片并提取关键词，形成可检索来源片段。",
        146: "图 6-1 V1 来源导入、切片与资源生成流程",
        147: "6.2 网络来源搜索与导入",
        148: "来源面板支持输入检索词搜索网页候选，并允许选择单条或批量导入。候选结果包含标题、链接、摘要、域名和搜索提供方，导入后与本地文件共享同一来源结构。",
        149: "正文抽取按优先级尝试 Jina Reader、Trafilatura 和 HTML 清洗。对于反爬、登录页或动态网页等无法获得完整正文的情况，系统保留搜索摘要并标记为 fallback，避免整个导入流程失败。",
        150: "导入后的文本按约 720 字符切片，并保留约 120 字符重叠；切片数量由正文长度决定。前端刷新工作区后在来源列表展示导入状态、摘要和来源详情。",
        151: "批量网页导入由后端限制并发并合并写入，降低 JSON 状态被并发覆盖的风险。前端使用 React 状态和 API 请求同步来源列表，不依赖 Server Actions 或 Zustand。",
        152: "图 6-2 网络来源搜索与候选导入界面",
        154: "6.3 云大学堂转写导入与片段播放",
        155: "系统可通过认证信息或 Cookie 建立云大学堂会话，查询课程并导入可获取的课堂转写。导入后的来源类型为 lecture，包含课程和课程片段元数据。",
        156: "课堂转写片段保存 start_time、end_time、start_seconds、end_seconds、video_url 和 source_url 等元数据。没有可播放视频时，系统仍展示来源片段和课程页面信息。",
        157: "当问答引用命中课堂转写片段时，聊天区将展示可打开的课堂引用卡片。",
        158: "引用卡片使用浏览器原生 video 元素或课程 iframe 打开内容，并根据 start_seconds 跳转到对应时间段播放；该能力依赖课程页面和视频地址的可访问性。",
        160: "图 6-3 云大学堂课程选择与转写导入",
        165: "图 6-4 课堂转写来源示例",
        166: "6.4 轻量 RAG 问答与流式输出",
        167: "用户提问后，Retrieval 模块从 ready 状态来源中检索关键词、全文和元数据匹配的片段，并对来源质量进行加权。Tutor 模块将相关片段与问题组成模型上下文，返回回答、来源标题、位置和片段摘要。后端使用 HTTP 流式响应向前端逐段返回文本；未命中直接来源时，会在回答中标明“以下为通用解释”。",
        168: "6.5 学习画像与学习建议",
        169: "学习画像包括知识基础描述、学习目标、薄弱点、资源偏好、正确率、学习节奏、易错模式、兴趣和下一步建议。聊天内容可补充兴趣信息，测验和闪卡反馈可更新薄弱点与资源偏好。",
        170: "闪卡复习记录“会 / 不会 / 跳过”。其中“不会”会将对应主题加入或前置到薄弱点；系统据此生成下一步复习建议。",
        171: "测验提交后会计算正确率，并将错题主题写入薄弱点列表。",
        172: "学习建议以“复习薄弱点 - 完成闪卡 - 进行测验 - 展开导图追问”等顺序组织，帮助学生建立可操作的学习路径。",
        173: "画像更新基于当前学习事件，不包含固定分值加减、遗忘曲线调度或医学/心理学意义上的疲劳预测。",
        175: "6.6 交互式思维导图",
        176: "思维导图由资源生成模块输出树状 JSON，前端以自研 SVG 连线和 HTML 节点布局呈现。用户可展开或收起分支、缩放画布、下载 Markdown，并点击节点向聊天区发起基于来源的追问。",
        177: "图 6-5 思维导图与学习反馈界面",
        178: "图 6-6 闪卡与测验反馈界面",
        181: "6.7 演示文稿生成与浏览器端导出",
        182: "Studio 可根据当前来源生成结构化演示文稿，包括封面、要点页、时间线、双栏比较、问答和总结等布局。生成内容由 Resource 模块组织为 JSON 数据，前端以分页预览呈现。",
        183: "演示文稿生成优先使用已导入来源与学习主题；模型不可用时，系统可使用来源关键词和预设结构生成基础讲义。",
        184: "前端支持查看单页、缩略图导航、播放式展开和 PDF 导出。",
        185: "PPTX 导出使用浏览器端 PptxGenJS。当前导出会将前端渲染后的页面作为图片写入 PPTX，适合展示和提交，不承诺导出文件中的文字可编辑。",
        186: "演示文稿导出由浏览器完成，不依赖后端 python-pptx 服务。",
        187: "图 6-7 演示文稿预览与导出界面",
        188: "第七章：测试与验证",
        189: "7.1 测试环境和方法",
        190: "开发与验证环境：Windows 本地开发环境，前端使用 Node.js 和 Next.js，后端使用 Python、FastAPI 与 requirements.txt 中列出的依赖。项目可通过 start-dev.bat 启动前后端服务；Docker 部署文件使用 Python 3.12 和 Node.js 22 作为容器运行时。",
        191: "验证方式：执行 python -m compileall backend\\app、前端 npm run lint 和 npm run build，并对 /health、来源上传、网页导入、聊天、资源生成、测验、闪卡和课堂转写流程进行人工联调。",
        192: "本项目未进行 GPU 压测、Android 原生打包测试或大规模并发压测，因此不在本文中声明固定性能指标。",
        193: "7.2 核心功能验证",
        194: "表 7-1 以当前 V1 实现为范围的核心功能验证用例",
        196: "7.3 非功能性与限制说明",
        197: "内容安全与可靠性：Safety 模块对少量明显不安全的关键词进行基础拦截；模型提示词要求基于来源回答并区分来源依据和通用解释。网页抓取与模型调用均有降级路径，但该原型不宣称达到生产级内容审核或高可用指标。",
        198: "性能说明：模型响应时延受模型服务、网络、来源长度和本机环境影响。当前验证重点为功能正确性、加载状态和降级流程，不给出未经独立压测的固定延迟、并发量或帧率结论。",
        199: "7.4 演示场景验证",
        200: "演示场景验证：使用计算机网络课程种子知识、本地课程文件、网页来源和可获取的课堂转写进行端到端演示，验证从来源导入、来源问答、Studio 生成到学习反馈的闭环。",
        201: "教育效果评估属于后续工作。本 V1 尚未开展具有统计学结论的学生对照实验，因此不在本文中使用学习成绩提升或时间节省等量化结论。",
        203: "第八章：开源合规与 AI Coding 协作说明",
        204: "8.1 依赖与许可",
        205: "项目使用的开源依赖和用途如下。第三方服务与后续规划技术会单独标注，避免与当前运行依赖混淆。",
        206: "1. Next.js、React、TypeScript、Tailwind CSS：前端工作台与样式系统，均遵循各自开源许可证。",
        207: "2. html-to-image、jsPDF、PptxGenJS、KaTeX：浏览器端页面导出、演示文稿导出和数学公式渲染。",
        208: "3. FastAPI、Uvicorn、Pydantic：后端 API、服务运行和数据模型校验。",
        209: "4. pdfplumber、python-docx、python-pptx、Pillow、Trafilatura：PDF、DOCX、PPTX、图片文件支持和网页正文抽取。",
        210: "5. LangGraph、PostgreSQL + pgvector、PaddleOCR、LiteLLM：当前作为 V2 可选演进方案，不属于 V1 运行依赖。",
        211: "8.2 AI Coding 协作说明",
        212: "本项目在需求梳理、界面迭代、代码检查和文档校对过程中使用了 AI Coding 工具进行辅助。AI 产出由项目成员审阅、修改和测试后再纳入项目。",
        213: "协作边界：AI 用于生成候选实现、解释报错、辅助重构和文档校对；架构取舍、功能验收、测试执行和最终提交由项目成员负责。",
        214: "本文不以 AI 工具替代人工测试，也不宣称未经保存的性能优化、测试覆盖率或第三方实验结果。",
        216: "项目成员对关键交互、接口返回和演示流程进行了人工复核，并通过后端编译检查、前端 lint 与 build 验证基础工程质量。",
        217: "后续将逐步补充自动化测试、性能基线和多用户持久化能力。",
        219: "8.3 总结与提交说明",
        220: "NetNote V1 已完成面向计算机网络学习的来源导入、轻量 RAG 问答、学习资源生成、学习反馈、网页资料补充和课堂转写联动等能力。项目以可演示、可维护和可扩展为目标，当前源码、运行说明与本文档保持一致；向量检索、OCR、显式智能体编排和大规模效果评估作为后续迭代方向。",
    }
    heading_styles = {
        110: "Heading 2", 130: "Heading 1", 132: "Heading 2", 134: "Heading 2", 138: "Heading 2", 140: "Heading 2", 143: "Heading 1",
        144: "Heading 2", 147: "Heading 2", 154: "Heading 2", 166: "Heading 2", 168: "Heading 2", 175: "Heading 2", 181: "Heading 2",
        188: "Heading 1", 189: "Heading 2", 193: "Heading 2", 196: "Heading 2", 199: "Heading 2", 203: "Heading 1", 204: "Heading 2", 211: "Heading 2", 219: "Heading 2",
    }
    for index, text in replacements.items():
        replace_paragraph(paragraphs[index], text, heading_styles.get(index))

    # The source document contains an empty Heading 2 anchor beside a floating
    # screenshot. Keep the anchor for layout, but exclude it from the TOC.
    if not paragraphs[119].text.strip():
        paragraphs[119].style = "Normal"

    for paragraph in paragraphs:
        if paragraph.text.strip() == "中国网络软件杯赛—大学生软件设计大赛（2026）":
            replace_paragraph(paragraph, "中国软件杯大学生软件设计大赛（2026）")

    # Restore the two architecture visuals as inline figures. The source used
    # floating legacy drawings, which Word can drop during an A4 reflow.
    insert_image_after(paragraphs[116], architecture_image((1440, 900)), Mm(160))
    insert_image_after(paragraphs[126], workflow_image((1440, 820)), Mm(160))

    tables = doc.tables
    replace_table(tables[0], ["序号", "修订原因/主要内容", "版本号", "编制人", "修订日期", "备注说明"], [
        ["1", "完成 V1 核心学习流程与界面", "V1.0", "项目组", "2026.06.22", "完成来源、问答和 Studio 原型"],
        ["2", "完善来源导入、网页抽取与学习反馈", "V1.1", "项目组", "2026.06.28", "形成来源到反馈闭环"],
        ["3", "接入云大学堂转写和演示文稿能力", "V1.2", "项目组", "2026.07.03", "补充课堂片段播放与导出"],
        ["4", "校准 V1 文档与源码说明", "V1.3", "项目组", "2026.07.16", "明确已实现与后续规划"],
    ])
    replace_table(tables[1], ["字段名", "运行时类型", "说明", "示例/取值", "备注"], [
        ["id", "str", "来源唯一标识", "source_xxx", "工作区内唯一"],
        ["title", "str", "来源标题或文件名", "计算机网络讲义.pdf", "前端展示"],
        ["kind", "Literal", "来源类别", "file/web/seed/lecture", "统一来源入口"],
        ["extraction_status", "Literal", "正文抽取质量", "complete/partial/fallback", "网页导入可降级"],
        ["chunks", "list[SourceChunk]", "可检索文本片段", "文本、位置、关键词、metadata", "不含 embedding"],
    ])
    replace_table(tables[2], ["字段名", "运行时类型", "说明", "典型值", "更新方式"], [
        ["knowledge_base", "str", "知识基础描述", "入门到进阶", "初始化/人工说明"],
        ["learning_goal", "str", "学习目标", "掌握计算机网络", "工作区配置"],
        ["weak_points", "list[str]", "薄弱知识点", "TCP 拥塞控制", "测验与闪卡反馈"],
        ["preferred_resources", "list[str]", "资源偏好", "来源问答/闪卡", "闪卡反馈"],
        ["accuracy_rate", "float", "最近测验正确率", "0.75", "测验提交"],
        ["learning_pace", "str", "学习节奏建议", "每次聚焦 2-3 点", "默认/建议"],
        ["error_patterns", "list[str]", "易错模式", "概念混淆", "学习记录"],
        ["interests", "list[str]", "兴趣主题", "HTTP/TCP/IP", "聊天关键词"],
        ["next_steps", "list[str]", "下一步建议", "复习-闪卡-测验", "Profile 模块"],
    ])
    replace_table(tables[3], ["字段名", "运行时类型", "说明", "示例", "备注"], [
        ["id", "str", "片段唯一标识", "chunk_xxx", "工作区内唯一"],
        ["source_id", "str", "所属来源 ID", "source_xxx", "关联来源"],
        ["source_title", "str", "来源标题", "计算机网络讲义", "引用展示"],
        ["text", "str", "片段正文", "切片后的课程文本", "检索证据"],
        ["location", "str", "来源位置", "第 2 页 / 00:12:10", "引用定位信息"],
        ["keywords", "list[str]", "提取关键词", "TCP/拥塞控制", "轻量检索"],
        ["metadata", "dict", "扩展元数据", "课堂时间、视频地址", "按来源类型扩展"],
    ])
    replace_table(tables[4], ["字段名", "运行时类型", "说明", "示例", "备注"], [
        ["kind", "str", "来源类型", "lecture", "课堂转写来源"],
        ["course_id", "str", "课程标识", "course_xxx", "课程信息"],
        ["week / section", "str", "周次与节次", "第 3 周 / 第 2 节", "可为空"],
        ["start_seconds", "int", "片段起始秒数", "1455", "播放器跳转"],
        ["end_seconds", "int", "片段结束秒数", "1490", "片段范围"],
        ["video_url / source_url", "str", "视频或课程页面地址", "https://...", "依赖平台可访问性"],
    ])
    replace_table(tables[5], ["字段名", "运行时类型", "说明", "示例", "备注"], [
        ["id", "str", "生成物唯一标识", "artifact_xxx", "工作区内唯一"],
        ["kind", "Literal", "生成物类型", "quiz/presentation", "支持 7 类资源"],
        ["title", "str", "生成物标题", "计算机网络测验", "前端展示"],
        ["status", "Literal", "生成状态", "generating/ready/failed", "任务反馈"],
        ["data", "dict", "结构化资源内容", "cards/questions/slides", "按类型解析"],
    ])
    replace_table(tables[6], ["字段名", "运行时类型", "说明", "示例", "备注"], [
        ["id", "str", "消息唯一标识", "message_xxx", "工作区内唯一"],
        ["role", "Literal", "消息角色", "user/assistant/summary", "对话展示"],
        ["content", "str", "Markdown 文本内容", "回答正文", "支持流式累积"],
        ["citations", "list[Citation]", "引用信息", "来源标题/位置/片段", "按检索结果生成"],
    ])
    replace_table(tables[7], ["用例 ID", "验证模块", "执行路径", "预期结果", "验证结果", "状态"], [
        ["TC-IMP-01", "文件来源导入", "上传 PDF/DOCX/PPTX/TXT/MD", "来源入库并可查看文本片段", "可完成导入与查看", "通过"],
        ["TC-WEB-02", "网页候选导入", "搜索并导入网页候选", "获取正文或摘要降级结果", "显示来源质量状态", "通过"],
        ["TC-RAG-03", "来源问答", "基于已导入来源提问", "返回回答与来源引用信息", "聊天区显示引用条目", "通过"],
        ["TC-STU-04", "Studio 资源", "生成闪卡、测验、导图或讲义", "生成结构化资源并打开详情", "资源可交互查看", "通过"],
        ["TC-FBK-05", "学习反馈", "提交测验或记录闪卡结果", "更新正确率、薄弱点和建议", "工作区反馈更新", "通过"],
        ["TC-YNU-06", "课堂转写", "连接后导入可获取课堂转写", "保存时间元数据并可打开引用", "依赖课程访问权限", "人工验证"],
        ["TC-PPT-07", "演示文稿导出", "生成并导出 PDF/PPTX", "浏览器下载导出文件", "PPTX 为图片版", "通过"],
    ])

    doc.save(OUTPUT)
    # Word keeps one legacy empty heading anchor beside a floating object. A
    # second pass removes it after save so it cannot create a blank TOC entry.
    final_doc = Document(OUTPUT)
    for paragraph in final_doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and not paragraph.text.strip():
            paragraph.style = "Normal"
    final_doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
